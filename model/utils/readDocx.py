from docx import Document

def readDocx(file_path):
    doc=Document(file_path)
    full_texts=[]
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_texts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_texts.append(cell.text)

    document_text='\n'.join(full_texts)
    return document_text